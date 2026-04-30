from docx import Document
import sys

doc = Document(r'C:\Users\高原\Desktop\Get笔记API文档.docx')
lines = []
for p in doc.paragraphs:
    lines.append(p.text)

# 也读取表格
for table in doc.tables:
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        lines.append(' | '.join(cells))
    lines.append('')

with open(r'C:\Users\高原\.stepclaw\workspace-main-2\temp\getnote-api.md', 'w', encoding='utf-8') as f:
    f.write('\n'.join(lines))

print('Done, lines:', len(lines))
